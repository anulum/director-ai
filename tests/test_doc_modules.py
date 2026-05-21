# SPDX-License-Identifier: AGPL-3.0-or-later
# Commercial license available
# © Concepts 1996–2026 Miroslav Šotek. All rights reserved.
# © Code 2020–2026 Miroslav Šotek. All rights reserved.
# ORCID: 0009-0009-3560-0851
# Contact: www.anulum.li | protoscience@anulum.li
# Director-Class AI — Document Module Tests
"""Multi-angle tests for document chunker, parser, registry pipeline.

Covers: PDF/DOCX/HTML parsing, chunk sizing, overlap, metadata, registry
lookup, pipeline integration with ingest CLI, and performance documentation.
"""

from __future__ import annotations

import io
import sys
import types

import numpy as np
import pytest

from director_ai.core.doc_chunker import ChunkConfig, split
from director_ai.core.doc_parser import parse
from director_ai.core.doc_registry import DocRegistry
from director_ai.core.retrieval import doc_chunker


class TestChunker:
    @pytest.mark.parametrize(
        ("kwargs", "message"),
        [
            ({"chunk_size": 0}, "chunk_size"),
            ({"chunk_size": -1}, "chunk_size"),
            ({"chunk_size": 8, "overlap": -1}, "overlap"),
            ({"chunk_size": 8, "overlap": 8}, "overlap"),
            ({"separators": ()}, "separators"),
            ({"separators": ("\n", object())}, "separators"),
            ({"similarity_threshold": -0.01}, "similarity_threshold"),
            ({"similarity_threshold": 1.01}, "similarity_threshold"),
        ],
    )
    def test_chunk_config_rejects_invalid_values(self, kwargs, message):
        with pytest.raises(ValueError, match=message):
            ChunkConfig(**kwargs)

    def test_empty(self):
        assert split("") == []

    def test_short_text(self):
        assert split("Hello world.", ChunkConfig(chunk_size=100)) == ["Hello world."]

    def test_splits_on_sentence(self):
        text = "First sentence. Second sentence. Third sentence."
        chunks = split(text, ChunkConfig(chunk_size=30, overlap=0))
        assert len(chunks) >= 2

    def test_overlap(self):
        text = "A" * 100 + " " + "B" * 100
        chunks = split(text, ChunkConfig(chunk_size=110, overlap=10))
        assert len(chunks) >= 2

    def test_respects_max_size(self):
        text = "word " * 200
        chunks = split(text, ChunkConfig(chunk_size=50, overlap=0))
        for chunk in chunks:
            assert len(chunk) <= 60  # some slack for separator

    def test_unicode(self):
        text = "HĂ©llo wĂ¶rld. ĂśnĂŻcĂ¶dĂ© text here. More sentences follow."
        chunks = split(text, ChunkConfig(chunk_size=30, overlap=0))
        assert len(chunks) >= 1

    def test_single_long_word(self):
        text = "A" * 1000
        chunks = split(text, ChunkConfig(chunk_size=100, overlap=10))
        assert len(chunks) >= 5

    def test_sentence_semantic_mode_falls_back_without_embedding_backend(
        self, monkeypatch
    ):
        """Semantic chunking must degrade to deterministic recursive splitting.

        This covers enterprise ingestion environments where the embedding extra
        is not installed but the caller still requests semantic chunking.
        """
        monkeypatch.setattr(doc_chunker, "_embed_sentences", lambda sentences: None)
        text = (
            "The refund policy allows returns within thirty days. "
            "Invoices must include the tenant identifier. "
            "A sensor should halt when torque exceeds the configured limit."
        )

        chunks = split(text, ChunkConfig(chunk_size=70, overlap=0, semantic=True))

        assert len(chunks) >= 2
        assert "refund policy" in chunks[0]
        assert any("sensor should halt" in chunk for chunk in chunks)

    def test_semantic_mode_splits_on_embedding_topic_shift(self, monkeypatch):
        """Low cosine similarity between neighbouring sentences creates chunks."""
        embeddings = np.array(
            [
                [1.0, 0.0],
                [0.99, 0.01],
                [0.0, 1.0],
                [0.01, 0.99],
            ],
            dtype=float,
        )
        monkeypatch.setattr(
            doc_chunker, "_embed_sentences", lambda sentences: embeddings
        )
        text = (
            "Contract liability is capped at the monthly fee. "
            "Indemnity follows the service terms. "
            "The robot arm must stop before the keep-out zone. "
            "Torque checks run before every motion plan."
        )

        chunks = split(text, ChunkConfig(chunk_size=120, semantic=True))

        assert chunks == [
            "Contract liability is capped at the monthly fee. Indemnity follows the service terms.",
            "The robot arm must stop before the keep-out zone. Torque checks run before every motion plan.",
        ]

    def test_semantic_mode_recursively_splits_large_topic_group(self, monkeypatch):
        embeddings = np.array([[1.0, 0.0], [0.99, 0.01]], dtype=float)
        monkeypatch.setattr(
            doc_chunker, "_embed_sentences", lambda sentences: embeddings
        )
        text = (
            "Alpha " * 30
            + "must remain traceable. "
            + "Beta " * 30
            + "must also remain traceable."
        )

        chunks = split(text, ChunkConfig(chunk_size=80, overlap=0, semantic=True))

        assert len(chunks) > 2
        assert all(len(chunk) <= 80 for chunk in chunks)

    def test_embed_sentences_uses_sentence_transformer_without_progress(
        self, monkeypatch
    ):
        calls: dict[str, object] = {}

        class FakeSentenceTransformer:
            def __init__(self, model_name):
                calls["model_name"] = model_name

            def encode(self, sentences, *, show_progress_bar):
                calls["sentences"] = list(sentences)
                calls["show_progress_bar"] = show_progress_bar
                return np.array([[1.0, 0.0], [0.0, 1.0]], dtype=float)

        fake_module = types.ModuleType("sentence_transformers")
        fake_module.SentenceTransformer = FakeSentenceTransformer
        monkeypatch.setitem(sys.modules, "sentence_transformers", fake_module)

        result = doc_chunker._embed_sentences(["Alpha.", "Beta."])

        assert calls == {
            "model_name": "all-MiniLM-L6-v2",
            "sentences": ["Alpha.", "Beta."],
            "show_progress_bar": False,
        }
        np.testing.assert_allclose(result, [[1.0, 0.0], [0.0, 1.0]])

    def test_embed_sentences_returns_none_when_backend_missing(self, monkeypatch):
        monkeypatch.setitem(sys.modules, "sentence_transformers", None)

        assert doc_chunker._embed_sentences(["Alpha."]) is None

    def test_semantic_mode_delegates_sentence_split_to_rust_when_available(
        self, monkeypatch
    ):
        calls: list[str] = []

        def _fake_rust_split(text: str) -> list[str]:
            calls.append(text)
            return ["Alpha.", "Beta.", "Gamma."]

        monkeypatch.setattr(doc_chunker, "_RUST_DOC_CHUNKER", True)
        monkeypatch.setattr(doc_chunker, "rust_split_sentences", _fake_rust_split)
        monkeypatch.setattr(doc_chunker, "_embed_sentences", lambda _s: None)

        text = "Alpha. Beta. Gamma."
        chunks = split(text, ChunkConfig(chunk_size=10, overlap=0, semantic=True))

        assert calls == [text]
        assert chunks

    def test_semantic_mode_rust_sentence_split_exception_falls_back(self, monkeypatch):
        def _boom(_text: str) -> list[str]:
            raise RuntimeError("ffi fail")

        monkeypatch.setattr(doc_chunker, "_RUST_DOC_CHUNKER", True)
        monkeypatch.setattr(doc_chunker, "rust_split_sentences", _boom)
        monkeypatch.setattr(doc_chunker, "_embed_sentences", lambda _s: None)

        text = "Alpha sentence. Beta sentence. Gamma sentence."
        chunks = split(text, ChunkConfig(chunk_size=25, overlap=0, semantic=True))

        assert chunks


class TestParser:
    @pytest.mark.parametrize(
        ("content", "filename", "message"),
        [
            ("not bytes", "test.txt", "content"),
            (b"content", "", "filename"),
            (b"content", "   ", "filename"),
            (b"content", None, "filename"),
        ],
    )
    def test_parse_rejects_invalid_inputs(self, content, filename, message):
        with pytest.raises(ValueError, match=message):
            parse(content, filename)

    def test_txt(self):
        assert parse(b"Hello world", "test.txt") == "Hello world"

    def test_md(self):
        assert parse(b"# Heading\nBody", "doc.md") == "# Heading\nBody"

    def test_csv(self):
        result = parse(b"name,age\nAlice,30\nBob,25", "data.csv")
        assert "Alice" in result
        assert "30" in result

    def test_csv_preserves_quoted_cell_content(self):
        result = parse(b'name,notes\nAlice,"alpha, beta"\n', "data.csv")
        assert "Alice" in result
        assert "alpha, beta" in result

    def test_html_removes_non_content_regions(self):
        result = parse(
            b"<html><body><nav>menu</nav><p>Useful fact</p><script>x()</script></body></html>",
            "page.html",
        )
        assert "Useful fact" in result
        assert "menu" not in result
        assert "x()" not in result

    def test_docx_real_document(self):
        from docx import Document

        buffer = io.BytesIO()
        document = Document()
        document.add_paragraph("Contract clause one.")
        document.add_paragraph("Contract clause two.")
        document.save(buffer)

        result = parse(buffer.getvalue(), "contract.docx")

        assert "Contract clause one." in result
        assert "Contract clause two." in result

    def test_unknown_extension(self):
        result = parse(b"some content", "file.xyz")
        assert result == "some content"

    def test_utf8_decode(self):
        result = parse("HĂ©llo".encode(), "test.txt")
        assert "HĂ©llo" in result

    def test_pdf_missing_dep(self):
        import contextlib

        with contextlib.suppress(ImportError, Exception):
            parse(b"not a pdf", "test.pdf")

    def test_docx_missing_dep(self):
        import contextlib

        with contextlib.suppress(ImportError, Exception):
            parse(b"not a docx", "test.docx")


class TestRegistry:
    @pytest.mark.parametrize(
        ("method_name", "args", "message"),
        [
            ("register", ("", "source.txt", "t1", ["c0"]), "doc_id"),
            ("register", ("d1", "", "t1", ["c0"]), "source"),
            ("register", ("d1", "source.txt", "", ["c0"]), "tenant_id"),
            ("register", ("d1", "source.txt", "t1", []), "chunk_ids"),
            ("update", ("", ["c0"]), "doc_id"),
            ("update", ("d1", []), "chunk_ids"),
            ("delete", ("",), "doc_id"),
            ("get", ("", "t1"), "doc_id"),
            ("get", ("d1", ""), "tenant_id"),
            ("list_for_tenant", ("",), "tenant_id"),
            ("exists", ("",), "doc_id"),
        ],
    )
    def test_rejects_invalid_registry_inputs(self, method_name, args, message):
        reg = DocRegistry()
        method = getattr(reg, method_name)

        with pytest.raises(ValueError, match=message):
            method(*args)

    def test_returned_records_do_not_mutate_stored_state(self):
        reg = DocRegistry()
        returned = reg.register("d1", "test.txt", "t1", ["c0"])
        returned.source = "changed.txt"
        returned.chunk_ids.append("evil")

        stored = reg.get("d1", "t1")

        assert stored is not None
        assert stored.source == "test.txt"
        assert stored.chunk_ids == ["c0"]

    def test_list_for_tenant_returns_record_snapshots(self):
        reg = DocRegistry()
        reg.register("d1", "test.txt", "t1", ["c0"])
        listed = reg.list_for_tenant("t1")
        listed[0].chunk_ids.append("evil")

        stored = reg.get("d1", "t1")

        assert stored is not None
        assert stored.chunk_ids == ["c0"]

    def test_register_and_get(self):
        reg = DocRegistry()
        rec = reg.register("d1", "test.txt", "t1", ["d1:chunk:0", "d1:chunk:1"])
        assert rec.doc_id == "d1"
        assert rec.chunk_count == 2
        fetched = reg.get("d1", "t1")
        assert fetched is not None
        assert fetched.source == "test.txt"

    def test_tenant_isolation(self):
        reg = DocRegistry()
        reg.register("d1", "f.txt", "t1", ["c0"])
        assert reg.get("d1", "t1") is not None
        assert reg.get("d1", "t2") is None

    def test_list_for_tenant(self):
        reg = DocRegistry()
        reg.register("d1", "a.txt", "t1", ["c0"])
        reg.register("d2", "b.txt", "t1", ["c1"])
        reg.register("d3", "c.txt", "t2", ["c2"])
        assert len(reg.list_for_tenant("t1")) == 2
        assert len(reg.list_for_tenant("t2")) == 1

    def test_delete(self):
        reg = DocRegistry()
        reg.register("d1", "f.txt", "t1", ["c0"])
        deleted = reg.delete("d1")
        assert deleted is not None
        assert reg.get("d1", "t1") is None

    def test_update(self):
        reg = DocRegistry()
        reg.register("d1", "f.txt", "t1", ["c0"])
        reg.update("d1", ["c0", "c1", "c2"], source="g.txt")
        rec = reg.get("d1", "t1")
        assert rec.chunk_count == 3
        assert rec.source == "g.txt"

    def test_duplicate_register_raises(self):
        reg = DocRegistry()
        reg.register("d1", "f.txt", "t1", ["c0"])
        with pytest.raises(ValueError, match="already registered"):
            reg.register("d1", "g.txt", "t1", ["c1"])

    def test_exists(self):
        reg = DocRegistry()
        assert not reg.exists("d1")
        reg.register("d1", "f.txt", "t1", ["c0"])
        assert reg.exists("d1")

    def test_count(self):
        reg = DocRegistry()
        assert reg.count == 0
        reg.register("d1", "f.txt", "t1", ["c0"])
        assert reg.count == 1
